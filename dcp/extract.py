"""Phase 4 — per-document text extraction and regex candidate-surfacing.

Two responsibilities, in service of the deep-read findings extractor:

1. **Text extraction with per-page caching.** Run pypdf once per document
   and cache the per-page text under `data/raw_text/<source>/<application_ref>/
   <sha[:16]>.pages.json`. The cache is the contract between the parsing
   stage and the LLM stage — either can be re-run independently. Re-running
   the parser is a no-op when the cache is present and matches the source
   SHA.

2. **Regex pre-pass over the per-page text.** Surface candidate sentences
   for the high-signal patterns the rubric calls out — `\\d+\\s*MW`,
   generator counts, fuel storage in hours/litres/tonnes. The output is
   `(document_sha, page_number, sentence)` tuples; the LLM (or the
   human-in-the-loop reading via Claude Code's Read tool) decides which
   are real findings and what structured shape they take.

Some PDFs are scanned-image-only (no text layer); pypdf returns empty
strings for those pages. Those pages fall back to **RapidOCR** (ONNX
PaddleOCR — non-generative CTC decoding, pinned weights), chosen over a
VLM deliberately: the OCR text is the substrate the verbatim-quote
verification gate checks against, and it must fail *noisily* on
illegible input (garbage characters) rather than *fluently* (a
generative model's plausible hallucination would let an invented quote
verify). Measured on the 2026-08 Barbour-round fetch: ~5% of documents
are scanned-only, mostly clean typed council forms. Per-page OCR use is
recorded in the cache payload (`ocr_pages`) for the audit trail. If
rapidocr isn't installed, behaviour degrades to the old contract
(empty pages, vision-capable reader handles the doc whole).

**Format is sniffed from content, not from the filename.** 2,082
documents in the corpus are not PDFs, and 255 of them arrived with no
usable extension at all (`.bin`) — a sample of 200 was 54 Word
documents, 42 Outlook messages, 16 spreadsheets and 18 scanned TIFFs.
Dispatching on the suffix would leave every one of those unread, so the
suffix is only a tiebreak: magic bytes decide, ZIP and OLE containers
are opened and identified by the streams they hold.

**Synthetic pagination is labelled as such.** Only a PDF has pages. A
`.docx` is split into ~3,000-character sections on paragraph
boundaries, a workbook into one section per worksheet, a deck into one
per slide — because page-scoring (dcp/deepread_select.py) needs
something page-shaped to score, and sending an 86-page supporting
statement as one blob defeats it. The cache records which kind of
division it is in `pagination`; anything other than `"pages"` means the
index is ours, not the document's, and must not be shown to a reader as
a page number.
"""

from __future__ import annotations

import dataclasses
import datetime as dt
import json
import logging
import re
from pathlib import Path
from typing import Iterable

log = logging.getLogger(__name__)

# Single cache root for parsed text. Mirrors the bytes layout under data/raw/
# so the (source, application_ref) prefix points to the same logical doc set.
RAW_TEXT_ROOT = Path("data/raw_text")


# ---------------------------------------------------------------------------
# Per-page text extraction + cache
# ---------------------------------------------------------------------------


@dataclasses.dataclass(frozen=True)
class ExtractedDoc:
    """Per-page text for one document, plus extraction metadata."""

    sha: str
    bytes_path: Path
    pages: list[str]  # index = 0-based page; pages[i] is the page's text
    engine: str
    extracted_at: str
    ocr_pages: tuple[int, ...] = ()  # 1-based pages whose text came from OCR
    # Whose division `pages` represents: "pages" for a real PDF page or a
    # scan frame, otherwise "sections" / "sheets" / "slides" — ours, and not
    # citable as a page number. Caches written before this field existed are
    # all PDFs, so absence reads as "pages".
    pagination: str = "pages"

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def native_pagination(self) -> bool:
        """True when an index into `pages` is a page number the reader can cite."""
        return self.pagination == "pages"


def cache_path_for(source: str, application_ref: str, sha: str) -> Path:
    """Cache file location for a given document SHA.

    `application_ref` keeps its slashes (matches the `data/raw/` layout); the
    file ends in `.pages.json` to make it obvious the payload is structured,
    not a flat dump.
    """
    return RAW_TEXT_ROOT / source / application_ref / f"{sha[:16]}.pages.json"


# Engines that mean "nothing read this document", written into caches by an
# extractor that had no loader for the format. They are not results, and a
# cache holding one must not stop a re-run or count as an empty document.
STALE_ENGINES = frozenset({"skipped", "unsupported"})


def cached_engine(cache: Path) -> str | None:
    """The engine recorded in a cache file, or None if unreadable/absent.

    Cheap enough to call per document only for the formats that might carry
    a stale engine — the payload holds the whole document text, and some are
    tens of megabytes.
    """
    try:
        return json.loads(cache.read_text()).get("engine")
    except (OSError, ValueError):
        return None


def is_stale_cache(cache: Path) -> bool:
    """True if this cache records a failure to load rather than a result."""
    return cache.exists() and cached_engine(cache) in STALE_ENGINES


def extract_pdf(bytes_path: Path) -> list[str]:
    """Pull text out of a PDF, one entry per page (empty string for image-only pages).

    Uses pypdf — already a project dep, fast enough for the top-100 corpus.
    Returns an empty list if the PDF can't be opened at all (encrypted, etc.).
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(bytes_path))
    except Exception as exc:
        log.warning("pypdf failed to open %s: %s", bytes_path, exc)
        return []
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            log.warning("pypdf failed on a page of %s: %s", bytes_path, exc)
            pages.append("")
    return pages


# A page whose stripped pypdf text is shorter than this is treated as
# image-only and sent to OCR. Drawings often carry a few characters of
# title-block text that pypdf does recover; the threshold keeps those in
# scope for OCR without re-processing genuinely text-layered pages.
OCR_MIN_CHARS = 25

_OCR_ENGINE = None  # lazy singleton; RapidOCR loads ~3 ONNX models on init


def _get_ocr_engine():
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        from rapidocr_onnxruntime import RapidOCR
        _OCR_ENGINE = RapidOCR()
    return _OCR_ENGINE


def _ocr_image_tesseract(pil_image, psm: str = "3") -> str:
    """OCR one PIL image via the tesseract binary (no pytesseract dep).

    Default engine: on clean typed English council documents it preserves
    inter-word spacing and reading order noticeably better than RapidOCR's
    English model (which concatenates words and emits full-width CJK
    punctuation — measured on the 2026-08 Fife/Enfield scanned samples).

    `psm` 3 segments the page assuming upright text; `psm` 1 adds
    orientation detection, which standalone photographs need — see
    `extract_image`.
    """
    import subprocess
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".png") as tmp:
        pil_image.save(tmp.name, format="PNG")
        proc = subprocess.run(
            ["tesseract", tmp.name, "stdout", "--psm", psm, "-l", "eng"],
            capture_output=True, text=True, timeout=120,
        )
    if proc.returncode != 0:
        raise RuntimeError(f"tesseract failed: {proc.stderr[:200]}")
    return proc.stdout


def _ocr_image_rapidocr(pil_image, psm: str = "3") -> str:  # noqa: ARG001
    """OCR one PIL image via RapidOCR (ONNX PaddleOCR). Alternative engine —
    stronger on low-quality scans and complex layouts, but weaker spacing on
    clean English text. `psm` is accepted and ignored: it is a tesseract
    concept, and the two engines share a call signature."""
    import numpy as np

    engine = _get_ocr_engine()
    result, _elapse = engine(np.asarray(pil_image.convert("RGB")))
    return "\n".join(item[1] for item in (result or []))


_OCR_BACKENDS = {"tesseract": _ocr_image_tesseract, "rapidocr": _ocr_image_rapidocr}


def ocr_pdf_pages(
    bytes_path: Path, page_indices: list[int], engine: str = "tesseract",
) -> dict[int, str]:
    """OCR specific (0-based) pages of a PDF via pypdfium2 + a non-generative
    OCR engine.

    Returns {page_index: text}. Both engines are deliberately non-generative
    (see module docstring): misreads surface as noisy characters, never as
    fluent invented prose, so the quote-verification substrate stays honest.
    Rendering at ~300 DPI (scale 300/72) — the sweet spot for typed council
    forms. Returns {} if the renderer or engine is unavailable.
    """
    ocr_image = _OCR_BACKENDS[engine]
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        log.warning("OCR unavailable (%s); leaving image-only pages empty", exc)
        return {}

    out: dict[int, str] = {}
    try:
        pdf = pdfium.PdfDocument(str(bytes_path))
    except Exception as exc:
        log.warning("pypdfium2 failed to open %s: %s", bytes_path, exc)
        return {}
    try:
        for idx in page_indices:
            try:
                page = pdf[idx]
                # 300 DPI for ordinary pages, but cap the rendered long edge
                # at ~4000 px: A0/A1 site plans at full 300 DPI produce
                # 100-megapixel bitmaps that stall tesseract for minutes and
                # add nothing for text recovery.
                long_edge_pts = max(page.get_size())
                scale = min(300 / 72, 4000 / long_edge_pts)
                bitmap = page.render(scale=scale)
                out[idx] = ocr_image(bitmap.to_pil())
            except Exception as exc:
                log.warning("OCR failed on page %d of %s: %s", idx + 1, bytes_path, exc)
    finally:
        pdf.close()
    return out


# ---------------------------------------------------------------------------
# Format sniffing
# ---------------------------------------------------------------------------

# Leading bytes that identify a format outright. Order matters: the two-byte
# BMP signature is last because it is short enough to collide.
_MAGIC: tuple[tuple[bytes, str], ...] = (
    (b"%PDF", "pdf"),
    (b"{\\rtf", "rtf"),
    (b"\xff\xd8\xff", "image"),
    (b"\x89PNG\r\n\x1a\n", "image"),
    (b"II*\x00", "image"),
    (b"MM\x00*", "image"),
    (b"GIF87a", "image"),
    (b"GIF89a", "image"),
    (b"BM", "image"),
)

# Local file header, empty archive, and spanned archive respectively — an
# empty .zip is eight bytes of end-of-central-directory and nothing else.
_ZIP_SIGS = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
_OLE_SIG = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# Byte-order marks, and the codec each one announces. UTF-16 text is full of
# NUL bytes, which would otherwise read as "binary, and unrecognised" — five
# Chorley documents are UTF-16LE HTML.
_BOMS: tuple[tuple[bytes, str], ...] = (
    (b"\xef\xbb\xbf", "utf-8-sig"),
    (b"\xff\xfe\x00\x00", "utf-32-le"),
    (b"\x00\x00\xfe\xff", "utf-32-be"),
    (b"\xff\xfe", "utf-16-le"),
    (b"\xfe\xff", "utf-16-be"),
)

# Suffixes trusted only when the content is inconclusive (plain text, mostly).
_SUFFIX_HINTS = {
    ".htm": "html", ".html": "html", ".xhtml": "html",
    ".eml": "eml", ".csv": "csv", ".txt": "text", ".xml": "html",
}


def _sniff_zip(bytes_path: Path) -> str:
    """Identify an Office/OpenDocument container by the members it holds."""
    import zipfile

    try:
        with zipfile.ZipFile(bytes_path) as zf:
            names = set(zf.namelist())
            mimetype = ""
            if "mimetype" in names:
                mimetype = zf.read("mimetype").decode("ascii", "replace").strip()
    except Exception as exc:
        log.warning("not a readable zip: %s (%s)", bytes_path, exc)
        return "unknown"
    if "word/document.xml" in names:
        return "docx"
    if "xl/workbook.xml" in names:
        return "xlsx"
    if "xl/workbook.bin" in names:
        # .xlsb: an OOXML container holding a *binary* workbook part, which
        # openpyxl cannot read. Looks like an .xlsx from the outside.
        return "xlsb"
    if any(n.startswith("ppt/slides/") for n in names):
        return "pptx"
    if mimetype.endswith("opendocument.text"):
        return "odt"
    if mimetype.endswith("opendocument.spreadsheet"):
        return "ods"
    if mimetype.endswith("opendocument.presentation"):
        return "odp"
    return "zip"


def _sniff_ole(bytes_path: Path) -> str:
    """Identify a legacy Compound File (Word 97, Excel 97, Outlook .msg).

    All three share the same container signature, so the streams inside are
    the only honest discriminator — and the three are not interchangeable:
    a `.msg` misread as a `.doc` yields nothing.
    """
    try:
        import olefile
    except ImportError:
        log.warning("olefile missing; cannot tell .doc from .msg in %s", bytes_path)
        return "unknown"
    try:
        with olefile.OleFileIO(bytes_path) as ole:
            streams = {"/".join(p) for p in ole.listdir()}
    except Exception as exc:
        log.warning("unreadable compound file %s: %s", bytes_path, exc)
        return "unknown"
    if any(s.startswith("__substg1.0_") for s in streams):
        return "msg"
    if "WordDocument" in streams:
        return "doc"
    if "Workbook" in streams or "Book" in streams:
        return "xls"
    if "PowerPoint Document" in streams:
        return "ppt"
    return "unknown"


def decode_text(data: bytes) -> str:
    """Decode bytes to text, honouring a byte-order mark when one is present.

    The mark itself is dropped. Only `utf-8-sig` consumes it; the UTF-16 and
    UTF-32 codecs leave a zero-width U+FEFF at the front, which would then
    sit inside the first quote the deep-read extracts and fail verbatim
    verification against the source.
    """
    for bom, codec in _BOMS:
        if data.startswith(bom):
            return data.decode(codec, "replace").lstrip("﻿")
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", "replace")


def _sniff_text(head: bytes, suffix: str) -> str:
    """Classify something that is plainly text: markup, an email, or prose."""
    text = decode_text(head)
    stripped = text.lstrip("﻿").lstrip()
    low = stripped[:400].lower()
    if low.startswith(("<!doctype html", "<html", "<?xml", "<rss", "<svg")):
        return "html"
    # RFC 822: a header block before the first blank line. Checking only the
    # first line is not enough — plenty of council letters open "From: ...".
    first_lines = stripped.splitlines()[:8]
    header_like = sum(1 for ln in first_lines
                      if re.match(r"^[A-Za-z\-]{2,40}:\s", ln))
    if header_like >= 3 and any(
            ln.lower().startswith(("received:", "message-id:", "mime-version:",
                                   "return-path:", "date:"))
            for ln in first_lines):
        return "eml"
    return _SUFFIX_HINTS.get(suffix, "text")


def sniff_format(bytes_path: Path) -> str:
    """What this document actually is, regardless of what it is called.

    Returns one of: pdf, docx, doc, rtf, xlsx, xls, ods, odt, pptx, ppt,
    odp, msg, eml, html, csv, text, image, zip, unknown.
    """
    try:
        with open(bytes_path, "rb") as fh:
            head = fh.read(8192)
    except OSError as exc:
        log.warning("cannot read %s: %s", bytes_path, exc)
        return "unknown"
    if not head:
        return "unknown"
    for sig, fmt in _MAGIC:
        if head.startswith(sig):
            return fmt
    if head.startswith(_ZIP_SIGS):
        return _sniff_zip(bytes_path)
    if head.startswith(_OLE_SIG):
        return _sniff_ole(bytes_path)
    has_bom = head.startswith(tuple(bom for bom, _codec in _BOMS))
    if b"\x00" in head[:1024] and not has_bom:
        return "unknown"  # binary, and none of the signatures matched
    return _sniff_text(head, bytes_path.suffix.lower())


# ---------------------------------------------------------------------------
# Non-PDF loaders
# ---------------------------------------------------------------------------

# Target size of a synthetic section, in characters. A dense A4 page of prose
# runs 2,500-3,500 characters, so sections land in the same range the page
# scorer was tuned against (dcp/deepread_select.py).
SECTION_CHARS = 3000


def paginate(blocks: Iterable[str], *, target: int = SECTION_CHARS) -> list[str]:
    """Group text blocks into page-sized sections, never splitting a block.

    A single oversized block (a table dumped as one paragraph) becomes its
    own section rather than being cut mid-sentence — the quote-verification
    gate compares against this text verbatim, so a cut would fail a quote
    that is really there.
    """
    sections: list[str] = []
    current: list[str] = []
    size = 0
    for block in blocks:
        block = block.rstrip()
        if not block.strip():
            continue
        if current and size + len(block) > target:
            sections.append("\n".join(current))
            current, size = [], 0
        current.append(block)
        size += len(block) + 1
    if current:
        sections.append("\n".join(current))
    return sections


def _strip_xml(data: bytes) -> str:
    """Text out of an OpenDocument / OOXML part, with paragraph breaks kept."""
    import html as html_mod

    text = data.decode("utf-8", "replace")
    text = re.sub(r"</(?:w:p|text:p|text:h|a:p|text:list-item|"
                  r"table:table-row|w:tr)>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return html_mod.unescape(text)


def extract_docx(bytes_path: Path) -> list[str]:
    """Body text of a .docx in document order, tables included.

    Iterates the body's own children rather than `paragraphs` then `tables`:
    a supporting statement's capacity figures are usually *in* a table, and
    reading all prose followed by all tables would separate every figure
    from the heading that says what it measures.
    """
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(str(bytes_path))
    blocks: list[str] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            blocks.append(Paragraph(child, doc).text)
        elif tag == "tbl":
            for row in Table(child, doc).rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                if para.text.strip():
                    blocks.append(para.text)
    return paginate(blocks)


def extract_doc(bytes_path: Path) -> list[str]:
    """Legacy Word 97-2003, via the macOS `textutil` binary.

    There is no maintained pure-Python reader for the binary Word format.
    `textutil` ships with macOS and both machines in this pipeline are Macs;
    elsewhere this degrades to no text (the document is then logged as
    not-extracted, not as empty — see scripts/deepread_run.py).
    """
    import shutil
    import subprocess

    if not shutil.which("textutil"):
        log.warning("textutil unavailable; cannot read legacy .doc %s", bytes_path)
        return []
    proc = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(bytes_path)],
        capture_output=True, timeout=180)
    if proc.returncode != 0:
        log.warning("textutil failed on %s: %s", bytes_path,
                    proc.stderr.decode("utf-8", "replace")[:200])
        return []
    text = proc.stdout.decode("utf-8", "replace")
    return paginate(text.splitlines())


def extract_rtf(bytes_path: Path) -> list[str]:
    from striprtf.striprtf import rtf_to_text

    raw = decode_text(bytes_path.read_bytes())
    return paginate(rtf_to_text(raw, errors="ignore").splitlines())


def extract_xlsx(bytes_path: Path) -> list[str]:
    """One section per worksheet — the workbook's own division, not ours.

    Handed the bytes rather than the path: openpyxl decides what it will
    open by looking at the *filename extension*, so a perfectly good
    workbook that arrived as `.bin` is refused outright.
    """
    import io

    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(bytes_path.read_bytes()),
                                    read_only=True, data_only=True)
    except Exception as exc:
        log.warning("openpyxl failed on %s: %s", bytes_path, exc)
        return []
    sheets: list[str] = []
    try:
        for ws in wb.worksheets:
            lines = [f"[sheet: {ws.title}]"]
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                if any(cells):
                    lines.append(" | ".join(cells).rstrip(" |"))
            sheets.append("\n".join(lines))
    finally:
        wb.close()
    return sheets


def extract_opendocument(bytes_path: Path) -> list[str]:
    """ODF text/spreadsheet/presentation — one part, chunked."""
    import zipfile

    try:
        with zipfile.ZipFile(bytes_path) as zf:
            content = zf.read("content.xml")
    except Exception as exc:
        log.warning("unreadable OpenDocument %s: %s", bytes_path, exc)
        return []
    return paginate(_strip_xml(content).splitlines())


def extract_pptx(bytes_path: Path) -> list[str]:
    """One section per slide, in slide order."""
    import zipfile

    try:
        with zipfile.ZipFile(bytes_path) as zf:
            names = sorted(
                (n for n in zf.namelist()
                 if re.fullmatch(r"ppt/slides/slide\d+\.xml", n)),
                key=lambda n: int(re.search(r"(\d+)", n.rsplit("/", 1)[1]).group(1)))
            return [_strip_xml(zf.read(n)).strip() for n in names]
    except Exception as exc:
        log.warning("unreadable pptx %s: %s", bytes_path, exc)
        return []


def extract_msg(bytes_path: Path) -> list[str]:
    """An Outlook message: envelope, body, attachment names.

    These are consultee responses — objections, and the technical challenges
    from statutory consultees — so the envelope matters as much as the body:
    who objected, and when.
    """
    import extract_msg as em

    try:
        msg = em.Message(str(bytes_path))
    except Exception as exc:
        log.warning("extract-msg failed on %s: %s", bytes_path, exc)
        return []
    try:
        head = [f"From: {msg.sender or ''}", f"To: {msg.to or ''}",
                f"Cc: {msg.cc or ''}", f"Date: {msg.date or ''}",
                f"Subject: {msg.subject or ''}"]
        names = [a.getFilename() for a in (msg.attachments or [])]
        if names:
            head.append("Attachments: " + ", ".join(n for n in names if n))
        body = msg.body or ""
    finally:
        msg.close()
    return paginate(["\n".join(head), *body.splitlines()])


def extract_eml(bytes_path: Path) -> list[str]:
    """RFC 822 mail via the standard library; prefers the plain-text part."""
    import email
    from email import policy

    try:
        msg = email.message_from_bytes(bytes_path.read_bytes(), policy=policy.default)
    except Exception as exc:
        log.warning("unparseable email %s: %s", bytes_path, exc)
        return []
    head = [f"{k}: {msg.get(k, '')}" for k in ("From", "To", "Cc", "Date", "Subject")]
    body = ""
    try:
        part = msg.get_body(preferencelist=("plain", "html"))
        if part is not None:
            raw = part.get_content()
            body = _html_to_text(raw) if part.get_content_subtype() == "html" else raw
    except Exception as exc:
        log.warning("no readable body in %s: %s", bytes_path, exc)
    return paginate(["\n".join(head), *body.splitlines()])


def _html_to_text(markup: str) -> str:
    from selectolax.parser import HTMLParser

    tree = HTMLParser(markup)
    for tag in tree.css("script, style"):
        tag.decompose()
    return tree.text(separator="\n")


def extract_html(bytes_path: Path) -> list[str]:
    markup = decode_text(bytes_path.read_bytes())
    return paginate(_html_to_text(markup).splitlines())


def extract_plaintext(bytes_path: Path) -> list[str]:
    return paginate(decode_text(bytes_path.read_bytes()).splitlines())


# Beyond this many frames a TIFF is a document scan large enough that OCR
# cost matters; the cap matches the per-document page budget elsewhere.
MAX_IMAGE_FRAMES = 200


def extract_image(bytes_path: Path, *, ocr: bool = True,
                  ocr_engine: str = "tesseract") -> list[str]:
    """OCR a standalone image. Multi-frame TIFFs give one section per frame.

    Councils' capture drivers emit multi-page TIFFs for scanned
    correspondence, so frames are genuine document pages here.

    OCR'd with orientation detection (`--psm 1`), unlike PDF pages. A
    standalone image in this corpus is often a photograph — a campaign
    banner, a site notice on a lamp post — and is as likely to be sideways
    or inverted as upright. Read upright-only, one Oxfordshire objection
    photograph returned `AUTHSCUOAXO TVUNY GNAIAC`, which is `DEFEND RURAL
    OXFORDSHIRE` backwards; with orientation detection it reads as written.
    """
    if not ocr:
        return []
    try:
        from PIL import Image
    except ImportError as exc:
        log.warning("Pillow missing; cannot OCR %s (%s)", bytes_path, exc)
        return []
    ocr_image = _OCR_BACKENDS[ocr_engine]
    try:
        img = Image.open(bytes_path)
    except Exception as exc:
        log.warning("unreadable image %s: %s", bytes_path, exc)
        return []
    pages: list[str] = []
    try:
        frames = min(getattr(img, "n_frames", 1), MAX_IMAGE_FRAMES)
        for idx in range(frames):
            try:
                img.seek(idx)
                pages.append(ocr_image(img.convert("RGB"), psm="1"))
            except Exception as exc:
                log.warning("OCR failed on frame %d of %s: %s", idx, bytes_path, exc)
                pages.append("")
    finally:
        img.close()
    return pages


def extract_zip(bytes_path: Path) -> list[str]:
    """Read the members of an archive, each through its own loader.

    These are not arbitrary archives: every one sampled is an exported
    email — a `header.txt` naming sender and subject, a `message.pdf`
    carrying the body, and the inline images. Consultee correspondence,
    in other words, which is where objections live.

    Members are read through `zf.read()` and written to a scratch file
    under a name of our own making; nothing is extracted to a path the
    archive chooses. Nested archives are not followed.
    """
    import tempfile
    import zipfile

    sections: list[str] = []
    try:
        with zipfile.ZipFile(bytes_path) as zf:
            members = [i for i in zf.infolist() if not i.is_dir()][:MAX_ZIP_MEMBERS]
            with tempfile.TemporaryDirectory() as tmpdir:
                for info in members:
                    try:
                        data = zf.read(info)
                    except Exception as exc:
                        log.warning("unreadable member %s of %s: %s",
                                    info.filename, bytes_path, exc)
                        continue
                    scratch = Path(tmpdir) / f"member{len(sections)}"
                    scratch.write_bytes(data)
                    fmt = sniff_format(scratch)
                    name = Path(info.filename).name
                    if fmt == "pdf":
                        pages = extract_pdf(scratch)
                    elif fmt in _LOADERS:
                        pages = _LOADERS[fmt][0](scratch)
                    else:
                        continue  # images and nested archives: not followed
                    body = "\n".join(p for p in pages if p.strip())
                    if body.strip():
                        sections.append(f"[{name}]\n{body}")
    except Exception as exc:
        log.warning("unreadable archive %s: %s", bytes_path, exc)
        return []
    return sections


# An archive with more members than this is not the email bundle this
# loader is for, and is left alone rather than read at unbounded cost.
MAX_ZIP_MEMBERS = 50

# Singular forms, for citing one of them. The plural is what the cache
# and documents.pagination store; this is what a reader is shown.
_PAGINATION_NOUN = {
    "pages": "page",
    "sections": "section",
    "sheets": "sheet",
    "slides": "slide",
}


def cite_page(page, pagination: str | None) -> str:
    """Name one `evidence_page` the way the source document divides itself.

    'page 4' for a PDF, 'section 4' for a Word file, 'sheet 2' for a
    workbook, 'slide 5' for a deck. This is the string a reporter uses to
    find a sentence before quoting it, so being wrong about it costs more
    than being vague: told "page 3" of a spreadsheet they open the file,
    find no page 3, and end up doubting the quote rather than the label.
    17,724 findings cite an index that is not a page.

    An unrecorded pagination yields a bare number rather than a guess.
    Most such documents are ordinary PDFs, but "most" is not a provenance
    claim, and this is the one field whose whole job is being checkable.
    """
    if page is None or page == "":
        return ""
    noun = _PAGINATION_NOUN.get(pagination or "")
    return f"{noun} {page}" if noun else str(page)


# format -> (loader, pagination label). The label says whose division the
# index represents; only "pages" may be shown to a reader as a page number.
_LOADERS = {
    "docx": (extract_docx, "sections"),
    "doc": (extract_doc, "sections"),
    "rtf": (extract_rtf, "sections"),
    "xlsx": (extract_xlsx, "sheets"),
    "ods": (extract_opendocument, "sections"),
    "odt": (extract_opendocument, "sections"),
    "odp": (extract_opendocument, "sections"),
    "pptx": (extract_pptx, "slides"),
    "msg": (extract_msg, "sections"),
    "eml": (extract_eml, "sections"),
    "html": (extract_html, "sections"),
    "csv": (extract_plaintext, "sections"),
    "text": (extract_plaintext, "sections"),
    "zip": (extract_zip, "sections"),
}

# Formats with no loader: recognised, but nothing can be read from them here.
# The binary pre-2007 Excel and PowerPoint formats and the binary `.xlsb`
# workbook, six documents corpus-wide between them.
UNSUPPORTED_FORMATS = {"xls", "xlsb", "ppt", "unknown"}


def extract_document(
    *,
    source: str,
    application_ref: str,
    sha: str,
    bytes_path: Path,
    force: bool = False,
    ocr: bool = True,
    ocr_engine: str = "tesseract",
) -> ExtractedDoc:
    """Return parsed per-page text for a document, using cache when present.

    Set `force=True` to bypass the cache (e.g. after upgrading the extraction
    engine). Pages with no usable text layer are OCR'd unless `ocr=False`;
    OCR'd page numbers are recorded in the cache for audit and the engine
    string becomes e.g. 'pypdf+tesseract'.
    """
    cache = cache_path_for(source, application_ref, sha)
    if cache.exists() and not force:
        payload = json.loads(cache.read_text())
        return ExtractedDoc(
            sha=payload["sha"],
            bytes_path=Path(payload["bytes_path"]),
            pages=payload["pages"],
            engine=payload["engine"],
            extracted_at=payload["extracted_at"],
            ocr_pages=tuple(payload.get("ocr_pages", [])),
            pagination=payload.get("pagination", "pages"),
        )

    ocr_pages: tuple[int, ...] = ()
    fmt = sniff_format(bytes_path)
    pagination = "sections"
    if fmt == "pdf":
        pages = extract_pdf(bytes_path)
        engine = "pypdf"
        pagination = "pages"
        if ocr:
            needs_ocr = [i for i, t in enumerate(pages)
                         if len(t.strip()) < OCR_MIN_CHARS]
            if needs_ocr:
                ocr_out = ocr_pdf_pages(bytes_path, needs_ocr, engine=ocr_engine)
                merged = {i: t for i, t in ocr_out.items() if t.strip()}
                for i, text in merged.items():
                    pages[i] = text
                if merged:
                    engine = f"pypdf+{ocr_engine}"
                    ocr_pages = tuple(sorted(i + 1 for i in merged))
    elif fmt == "image":
        pages = extract_image(bytes_path, ocr=ocr, ocr_engine=ocr_engine)
        engine = ocr_engine if pages else "unsupported"
        pagination = "pages"  # a scan's frames are the document's own pages
        ocr_pages = tuple(range(1, len(pages) + 1))
    elif fmt in _LOADERS:
        loader, pagination = _LOADERS[fmt]
        pages = loader(bytes_path)
        engine = fmt
        if not any(p.strip() for p in pages):
            # A loader that returned nothing is far more likely to have
            # failed than to have been handed a genuinely wordless
            # document — every loader here catches its own exceptions and
            # returns `[]`, so a broken library, a missing binary or a
            # corrupt file all arrive looking like an empty file. Caching
            # that would make a transient failure permanent, which is the
            # whole bug this dispatch was written to fix. Left uncached, it
            # is simply retried. PDFs keep the older contract: an
            # image-only page really is empty until OCR runs.
            log.info("%s loader returned nothing for %s; leaving uncached.",
                     fmt, bytes_path)
            return ExtractedDoc(
                sha=sha,
                bytes_path=bytes_path,
                pages=[],
                engine="unsupported",
                extracted_at=dt.datetime.now(dt.UTC).isoformat(timespec="seconds"),
                pagination=fmt,
            )
    else:
        # Recognised, but nothing here can read it. **Write no cache** — the
        # absence is what makes a later run retry, and it is what makes the
        # deep-read log `not_extracted` rather than `no_text`. An earlier
        # version wrote an empty cache here despite a comment promising it
        # did not, which recorded "we have no loader for this" as "this
        # document contains no words" and made the miss permanent.
        log.info("No loader for %s (sniffed=%s); leaving uncached.", bytes_path, fmt)
        return ExtractedDoc(
            sha=sha,
            bytes_path=bytes_path,
            pages=[],
            engine="unsupported",
            extracted_at=dt.datetime.now(dt.UTC).isoformat(timespec="seconds"),
            pagination=fmt,
        )

    doc = ExtractedDoc(
        sha=sha,
        bytes_path=bytes_path,
        pages=pages,
        engine=engine,
        extracted_at=dt.datetime.now(dt.UTC).isoformat(timespec="seconds"),
        ocr_pages=ocr_pages,
        pagination=pagination,
    )
    cache.parent.mkdir(parents=True, exist_ok=True)
    cache.write_text(json.dumps({
        "sha": doc.sha,
        "bytes_path": str(doc.bytes_path),
        "pages": doc.pages,
        "engine": doc.engine,
        "extracted_at": doc.extracted_at,
        "ocr_pages": list(doc.ocr_pages),
        "pagination": doc.pagination,
    }, ensure_ascii=False))
    return doc


# ---------------------------------------------------------------------------
# Regex pre-pass — high-signal patterns
# ---------------------------------------------------------------------------


# Capacity expressions. Captures the numeric magnitude and unit; tolerates a
# space or hyphen between number and unit, and optional decimals. The leading
# (?<![\w.]) and trailing (?![\w]) guards keep us from matching inside larger
# identifiers (e.g. "100MWh" still matches because the suffix may include 'h',
# but we don't want "FOO1MW1" to trigger).
CAPACITY_REGEX = re.compile(
    r"(?<![\w.])(\d{1,4}(?:\.\d+)?)\s*-?\s*(MW|kVA|kW|MVA)\b",
    re.IGNORECASE,
)

# Generator counts: "14 generators", "12 × diesel generators", "twenty-five
# gas reciprocating engines". We catch the digit form here; the LLM step
# handles spelled-out numbers if they show up.
GENERATOR_COUNT_REGEX = re.compile(
    r"(?<![\w.])(\d{1,3})\s*(?:×|x|\*)?\s*"
    r"(?:new\s+|proposed\s+|standby\s+|emergency\s+|backup\s+|back-up\s+|diesel\s+|gas\s+)*"
    r"(?:reciprocating\s+)?(?:engine\s+)?generators?\b",
    re.IGNORECASE,
)

# Fuel storage expressions — hours of run-time, litres, tonnes.
FUEL_STORAGE_REGEX = re.compile(
    r"(?<![\w.])(\d{1,6}(?:,\d{3})*(?:\.\d+)?)\s*"
    r"(hour|hours|hr|hrs|litres?|l|tonnes?|t)\s+"
    r"(?:of\s+)?(?:diesel|gas|fuel|LPG|propane)\b",
    re.IGNORECASE,
)

# Sentence splitter — naive but sufficient for surfacing candidates. We
# don't need perfect sentence boundaries; just enough context for a human
# (or LLM) to read the candidate phrase and assess it.
_SENT_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[A-Z])|\n{2,}")


@dataclasses.dataclass(frozen=True)
class Candidate:
    """One regex hit, with enough context to feed the LLM."""

    doc_sha: str
    page: int  # 1-based for display
    pattern: str  # 'capacity' | 'generator_count' | 'fuel_storage'
    match_text: str  # the regex match itself (e.g. "21MW", "14 gas reciprocating engine generators")
    sentence: str  # surrounding sentence

    def to_dict(self) -> dict:
        return dataclasses.asdict(self)


PATTERNS: dict[str, re.Pattern[str]] = {
    "capacity": CAPACITY_REGEX,
    "generator_count": GENERATOR_COUNT_REGEX,
    "fuel_storage": FUEL_STORAGE_REGEX,
}


def find_candidates(doc: ExtractedDoc) -> list[Candidate]:
    """Run every pattern against every page; return candidate hits with context.

    Each match yields one candidate with the enclosing sentence preserved.
    Duplicate sentences (same page, same pattern, same sentence text) are
    de-duplicated so the LLM-feed isn't padded with repeats.
    """
    seen: set[tuple[int, str, str]] = set()
    hits: list[Candidate] = []
    for page_idx, page_text in enumerate(doc.pages):
        if not page_text:
            continue
        sentences = _SENT_SPLIT.split(page_text)
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            for label, pattern in PATTERNS.items():
                for m in pattern.finditer(sentence):
                    key = (page_idx + 1, label, sentence)
                    if key in seen:
                        continue
                    seen.add(key)
                    hits.append(Candidate(
                        doc_sha=doc.sha,
                        page=page_idx + 1,
                        pattern=label,
                        match_text=m.group(0),
                        sentence=sentence,
                    ))
    return hits


def candidates_for_application(
    *,
    source: str,
    application_ref: str,
    documents: Iterable[tuple[str, Path]],
    force_extract: bool = False,
) -> list[Candidate]:
    """End-to-end pre-pass for one application: extract every doc, run regex.

    `documents` is an iterable of `(sha, bytes_path)` pairs (typically from
    the `documents` table or a manifest file). Returns the flat list of
    candidates across all docs, ordered first by doc-of-appearance then by
    page number.
    """
    out: list[Candidate] = []
    for sha, bytes_path in documents:
        doc = extract_document(
            source=source, application_ref=application_ref,
            sha=sha, bytes_path=bytes_path, force=force_extract,
        )
        out.extend(find_candidates(doc))
    return out
